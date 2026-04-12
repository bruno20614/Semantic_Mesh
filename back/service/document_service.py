import os
from datetime import datetime
from service.orm import Document, DocumentFile, DocumentContent

# Caminho padrão do Tesseract no Windows
try:
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
except ImportError:
    pass

IMAGE_EXTENSIONS = ('.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.webp')


def _extract_text_ocr(file_path):
    """Extrai texto de imagem via pytesseract (OCR)."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang='por+eng')
        return text.strip()
    except ImportError:
        return "Erro: pytesseract ou Pillow não instalados. Execute: pip install pytesseract Pillow"
    except Exception as e:
        return f"Erro ao aplicar OCR: {str(e)}"


def _extract_text_from_pdf_with_ocr(file_path):
    """Tenta extrair texto do PDF. Se vazio (PDF escaneado), aplica OCR página a página."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or '' for page in reader.pages).strip()
        if text:
            return text
        # PDF escaneado — converte páginas em imagem e aplica OCR
        try:
            from pdf2image import convert_from_path
            import pytesseract
            pages = convert_from_path(file_path, dpi=300)
            ocr_text = "\n".join(
                pytesseract.image_to_string(page, lang='por+eng') for page in pages
            )
            return ocr_text.strip() or "Nenhum texto detectado no PDF."
        except ImportError:
            return "PDF escaneado detectado. Instale pdf2image e pytesseract para extrair texto de PDFs escaneados."
        except Exception as e:
            return f"Erro ao aplicar OCR no PDF: {str(e)}"
    except Exception as e:
        return f"Erro ao extrair PDF: {str(e)}"


def save_document_with_content(db, org_id, file):
    file_location = os.path.join("..", "front", "uploads", str(org_id))
    os.makedirs(file_location, exist_ok=True)
    file_path = os.path.join(file_location, file.filename)
    with open(file_path, "wb") as f:
        f.write(file.file.read())
    doc = Document(organization_id=org_id, filename=file.filename, file_type=file.content_type, created_at=datetime.utcnow())
    db.add(doc)
    db.commit()
    db.refresh(doc)
    doc_file = DocumentFile(document_id=doc.id, file_path=file_path, file_hash="", uploaded_at=datetime.utcnow())
    db.add(doc_file)
    # --- Extração de texto ---
    raw_text = ""
    fname = file.filename.lower()
    try:
        if fname.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as txtf:
                raw_text = txtf.read()
        elif fname.endswith('.pdf'):
            raw_text = _extract_text_from_pdf_with_ocr(file_path)
        elif fname.endswith('.docx'):
            try:
                import docx
                docx_file = docx.Document(file_path)
                raw_text = "\n".join([p.text for p in docx_file.paragraphs])
            except Exception as e:
                raw_text = f"Erro ao extrair DOCX: {str(e)}"
        elif fname.endswith(IMAGE_EXTENSIONS):
            raw_text = _extract_text_ocr(file_path)
    except Exception as e:
        raw_text = f"Erro ao extrair texto: {str(e)}"
    if raw_text:
        doc_content = DocumentContent(document_id=doc.id, raw_text=raw_text)
        db.add(doc_content)
    db.commit()
    return doc


# Função para remover documento e arquivos

def remove_document(db, org_id, doc_id):
    doc = db.query(Document).filter(Document.id == doc_id, Document.organization_id == org_id).first()
    if not doc:
        return False
    doc_file = db.query(DocumentFile).filter(DocumentFile.document_id == doc_id).first()
    if doc_file:
        try:
            if os.path.exists(doc_file.file_path):
                os.remove(doc_file.file_path)
        except Exception:
            pass
        db.delete(doc_file)
    doc_content = db.query(DocumentContent).filter(DocumentContent.document_id == doc_id).first()
    if doc_content:
        db.delete(doc_content)
    db.delete(doc)
    db.commit()
    return True
