from fastapi import APIRouter, HTTPException, Depends, Request, UploadFile, File, Form
from interface.organization_model import OrganizationCreate, OrganizationOut
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from service.organization_service import (
    create_organization_service, list_organizations_service, update_organization_service, delete_organization_service,
    add_user_to_organization_service, remove_user_from_organization_service
)
from service.user_service import get_user_by_email_service
from uuid import UUID
from typing import List

router = APIRouter()
templates = Jinja2Templates(directory="../front/templates")

@router.post("/organization", response_model=OrganizationOut)
def create_organization(org: OrganizationCreate):
    result = create_organization_service(org.name)
    if not result:
        raise HTTPException(status_code=400, detail="Erro ao criar organização.")
    return result

@router.get("/organizations", response_class=HTMLResponse)
def organizations_html(request: Request):
    if 'user_id' not in request.session:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/", status_code=302)
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, UserOrganization
        db = SessionLocal()
    from service.orm import Organization, UserOrganization
    user_id = request.session['user_id']
    org_ids = [uo.organization_id for uo in db.query(UserOrganization).filter(UserOrganization.user_id == user_id).all()]
    orgs = db.query(Organization).filter(Organization.id.in_(org_ids)).all()
    db.close()
    return templates.TemplateResponse("organizations.html", {"request": request, "organizations": orgs})

@router.get("/organizations/create", response_class=HTMLResponse)
def organizations_create(request: Request, message: str = None, success: bool = None):
    return templates.TemplateResponse("organizations_create.html", {"request": request, "message": message, "success": success})

@router.post("/organizations/create", response_class=HTMLResponse)
def organizations_create_post(request: Request, name: str = Form(...)):
    result = create_organization_service(name)
    if not result:
        return templates.TemplateResponse("organizations_create.html", {"request": request, "message": "Erro ao criar organização.", "success": False})
    # Vincula o usuário criador à organização
    if 'user_id' in request.session:
        from service.orm import SessionLocal, UserOrganization
        db = SessionLocal()
        user_org = UserOrganization(user_id=request.session['user_id'], organization_id=result.id, role="owner")
        db.add(user_org)
        db.commit()
        db.close()
    return templates.TemplateResponse("organizations_create.html", {"request": request, "message": "Organização criada com sucesso!", "success": True})

@router.get("/organizations/api", response_model=List[OrganizationOut])
def list_organizations():
    return list_organizations_service()

@router.put("/organization/{org_id}", response_model=OrganizationOut)
def update_organization(org_id: UUID, org: OrganizationCreate):
    result = update_organization_service(org_id, org.name)
    if not result:
        raise HTTPException(status_code=404, detail="Organização não encontrada.")
    return result

@router.post("/organizations/{org_id}/delete")
def delete_organization_post(org_id: UUID):
    success = delete_organization_service(org_id)
    if not success:
        raise HTTPException(status_code=404, detail="Organização não encontrada.")
    from fastapi.responses import RedirectResponse
    from fastapi.responses import RedirectResponse
    return RedirectResponse("/organizations", status_code=303)

@router.delete("/organization/{org_id}")
def delete_organization(org_id: UUID):
    success = delete_organization_service(org_id)
    if not success:
        raise HTTPException(status_code=404, detail="Organização não encontrada.")
    return {"msg": "Organização deletada com sucesso!"}

# Adicionar usuário à organização
from fastapi import Form
@router.post("/organizations/{org_id}/add_user", response_class=HTMLResponse)
def add_user_to_organization(request: Request, org_id: UUID, email: str = Form(...)):
    user = get_user_by_email_service(email)
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado.")
    result = add_user_to_organization_service(user['id'], org_id)
    if not result:
        raise HTTPException(status_code=400, detail="Erro ao adicionar usuário à organização.")
    # Após adicionar, redireciona para a tela de edição da organização com mensagem de sucesso
    from fastapi.templating import Jinja2Templates
    templates = Jinja2Templates(directory="../front/templates")
    # Buscar dados atualizados da organização
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, Document, DocumentFile
        db = SessionLocal()
    from service.orm import Organization, Document, DocumentFile, UserOrganization, User
    org = db.query(Organization).filter(Organization.id == org_id).first()
    documents = []
    if org:
        docs = db.query(Document).filter(Document.organization_id == org_id).all()
        for doc in docs:
            doc_file = db.query(DocumentFile).filter(DocumentFile.document_id == doc.id).first()
            documents.append({
                "id": str(doc.id),
                "name": doc.filename,
                "uploaded_at": doc_file.uploaded_at if doc_file else doc.created_at
            })
    user_orgs = db.query(UserOrganization).filter(UserOrganization.organization_id == org_id).all()
    users = [db.query(User).filter(User.id == uo.user_id).first() for uo in user_orgs]
    db.close()
    return templates.TemplateResponse("organization_edit.html", {"request": request, "org": org, "documents": documents, "users": users, "message": "Usuário adicionado com sucesso!", "success": True})

# Remover usuário da organização
@router.post("/organizations/{org_id}/remove_user", response_class=HTMLResponse)
def remove_user_from_organization(request: Request, org_id: UUID, user_id: str = Form(...)):
    result = remove_user_from_organization_service(user_id, org_id)
    if not result:
        raise HTTPException(status_code=400, detail="Erro ao remover usuário da organização.")
    # Após remover, redireciona para a tela de edição da organização com mensagem de sucesso
    from fastapi.templating import Jinja2Templates
    templates = Jinja2Templates(directory="../front/templates")
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, Document, DocumentFile
        db = SessionLocal()
    from service.orm import Organization, Document, DocumentFile, UserOrganization, User
    org = db.query(Organization).filter(Organization.id == org_id).first()
    documents = []
    if org:
        docs = db.query(Document).filter(Document.organization_id == org_id).all()
        for doc in docs:
            doc_file = db.query(DocumentFile).filter(DocumentFile.document_id == doc.id).first()
            documents.append({
                "id": str(doc.id),
                "name": doc.filename,
                "uploaded_at": doc_file.uploaded_at if doc_file else doc.created_at
            })
    user_orgs = db.query(UserOrganization).filter(UserOrganization.organization_id == org_id).all()
    users = [db.query(User).filter(User.id == uo.user_id).first() for uo in user_orgs]
    db.close()
    return templates.TemplateResponse("organization_edit.html", {"request": request, "org": org, "documents": documents, "users": users, "message": "Usuário removido com sucesso!", "success": True})

# Tela de edição da organização (nome + documentos)
@router.get("/organizations/{org_id}/edit", response_class=HTMLResponse)
def organization_edit(request: Request, org_id: UUID, error: str = None):
    if 'user_id' not in request.session:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/", status_code=302)
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, Document, DocumentFile, UserOrganization
        db = SessionLocal()
    from service.orm import Organization, Document, DocumentFile, UserOrganization, User
    user_id = request.session['user_id']
    # Verifica se o usuário pertence à organização
    user_org = db.query(UserOrganization).filter(UserOrganization.user_id == user_id, UserOrganization.organization_id == org_id).first()
    if not user_org:
        db.close()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/organizations", status_code=302)
    org = db.query(Organization).filter(Organization.id == org_id).first()
    documents = []
    if org:
        docs = db.query(Document).filter(Document.organization_id == org_id).all()
        for doc in docs:
            doc_file = db.query(DocumentFile).filter(DocumentFile.document_id == doc.id).first()
            documents.append({
                "id": str(doc.id),
                "name": doc.filename,
                "uploaded_at": doc_file.uploaded_at if doc_file else doc.created_at
            })
    user_orgs = db.query(UserOrganization).filter(UserOrganization.organization_id == org_id).all()
    users = [db.query(User).filter(User.id == uo.user_id).first() for uo in user_orgs]
    db.close()
    return templates.TemplateResponse("organization_edit.html", {"request": request, "org": org, "documents": documents, "users": users, "error": error})

@router.post("/organizations/{org_id}/edit", response_class=HTMLResponse)
def organization_edit_post(request: Request, org_id: UUID, name: str = Form(...)):
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization
        db = SessionLocal()
    from service.orm import Organization
    org = db.query(Organization).filter(Organization.id == org_id).first()
    if not org:
        db.close()
        return templates.TemplateResponse("organization_edit.html", {"request": request, "org": None, "documents": [], "error": "Organização não encontrada."})
    org.name = name
    db.commit()
    db.close()
    # Após salvar, recarrega a tela com sucesso
    return templates.TemplateResponse("organization_edit.html", {"request": request, "org": org, "documents": [], "error": None})

# Tela de usuários da organização
@router.get("/organizations/{org_id}/users", response_class=HTMLResponse)
def organization_users(request: Request, org_id: UUID):
    if 'user_id' not in request.session:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/", status_code=302)
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, User, UserOrganization
        db = SessionLocal()
    from service.orm import Organization, User, UserOrganization
    user_id = request.session['user_id']
    user_org = db.query(UserOrganization).filter(UserOrganization.user_id == user_id, UserOrganization.organization_id == org_id).first()
    if not user_org:
        db.close()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/organizations", status_code=302)
    org = db.query(Organization).filter(Organization.id == org_id).first()
    user_orgs = db.query(UserOrganization).filter(UserOrganization.organization_id == org_id).all()
    users = [db.query(User).filter(User.id == uo.user_id).first() for uo in user_orgs]
    db.close()
    return templates.TemplateResponse("organization_users.html", {"request": request, "org": org, "users": users})

# Tela de documentos da organização (início CRUD)
@router.get("/organizations/{org_id}/documents", response_class=HTMLResponse)
def organization_documents(request: Request, org_id: UUID):
    if 'user_id' not in request.session:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/", status_code=302)
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, Document, DocumentFile, UserOrganization
        db = SessionLocal()
    from service.orm import Organization, Document, DocumentFile, UserOrganization
    user_id = request.session['user_id']
    user_org = db.query(UserOrganization).filter(UserOrganization.user_id == user_id, UserOrganization.organization_id == org_id).first()
    if not user_org:
        db.close()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/organizations", status_code=302)
    org = db.query(Organization).filter(Organization.id == org_id).first()
    documents = []
    if org:
        docs = db.query(Document).filter(Document.organization_id == org_id).all()
        for doc in docs:
            doc_file = db.query(DocumentFile).filter(DocumentFile.document_id == doc.id).first()
            documents.append({
                "id": str(doc.id),
                "name": doc.filename,
                "uploaded_at": doc_file.uploaded_at if doc_file else doc.created_at
            })
    db.close()
    return templates.TemplateResponse("organization_documents.html", {"request": request, "org": org, "documents": documents})

# Tela para adicionar documento (template futuro)
from fastapi import UploadFile, File, Form
from fastapi import Request
from fastapi.responses import RedirectResponse

@router.get("/organizations/{org_id}/documents/add", response_class=HTMLResponse)
def organization_documents_add(request: Request, org_id: UUID, message: str = None, success: bool = None):
    if 'user_id' not in request.session:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/", status_code=302)
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, UserOrganization
        db = SessionLocal()
    from service.orm import Organization, UserOrganization
    user_id = request.session['user_id']
    user_org = db.query(UserOrganization).filter(UserOrganization.user_id == user_id, UserOrganization.organization_id == org_id).first()
    if not user_org:
        db.close()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/organizations", status_code=302)
    org = db.query(Organization).filter(Organization.id == org_id).first()
    db.close()
    return templates.TemplateResponse("organization_documents_add.html", {"request": request, "org": org, "message": message, "success": success})

@router.post("/organizations/{org_id}/documents/add", response_class=HTMLResponse)
def organization_documents_add_post(request: Request, org_id: UUID, file: UploadFile = File(...)):
    if 'user_id' not in request.session:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/", status_code=302)
    import os
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Organization, Document, DocumentFile, UserOrganization
        db = SessionLocal()
    from service.orm import Organization, Document, DocumentFile, UserOrganization
    user_id = request.session['user_id']
    user_org = db.query(UserOrganization).filter(UserOrganization.user_id == user_id, UserOrganization.organization_id == org_id).first()
    if not user_org:
        db.close()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/organizations", status_code=302)
    org = db.query(Organization).filter(Organization.id == org_id).first()
    if not org:
        db.close()
        return templates.TemplateResponse("organization_documents_add.html", {"request": request, "org": None, "message": "Organização não encontrada.", "success": False})
    try:
        upload_dir = os.path.join("..", "front", "uploads", str(org_id))
        os.makedirs(upload_dir, exist_ok=True)
        file_location = os.path.join(upload_dir, file.filename)
        with open(file_location, "wb") as f:
            f.write(file.file.read())
        from datetime import datetime
        doc = Document(organization_id=org_id, filename=file.filename, file_type=file.content_type, created_at=datetime.utcnow())
        db.add(doc)
        db.commit()
        db.refresh(doc)
        doc_file = DocumentFile(document_id=doc.id, file_path=file_location, file_hash="", uploaded_at=datetime.utcnow())
        db.add(doc_file)
        # --- Extração de texto do documento ---
        raw_text = ""
        try:
            if file.filename.lower().endswith('.txt'):
                with open(file_location, 'r', encoding='utf-8') as txtf:
                    raw_text = txtf.read()
            elif file.filename.lower().endswith('.pdf'):
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(file_location)
                    raw_text = "\n".join(page.extract_text() or '' for page in reader.pages)
                except Exception as e:
                    raw_text = f"Erro ao extrair PDF: {str(e)}"
            elif file.filename.lower().endswith('.docx'):
                try:
                    import docx
                    docx_file = docx.Document(file_location)
                    raw_text = "\n".join([p.text for p in docx_file.paragraphs])
                except Exception as e:
                    raw_text = f"Erro ao extrair DOCX: {str(e)}"
        except Exception as e:
            raw_text = f"Erro ao extrair texto: {str(e)}"
        if raw_text:
            from service.orm import DocumentContent
            doc_content = DocumentContent(document_id=doc.id, raw_text=raw_text)
            db.add(doc_content)
        db.commit()
        db.close()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(f"/organizations/{org_id}/documents", status_code=303)
    except Exception as e:
        db.close()
        # Recarregar a organização da sessão para evitar DetachedInstanceError
        from service.orm import SessionLocal, Organization
        db = SessionLocal()
        org_fresh = db.query(Organization).filter(Organization.id == org_id).first()
        db.close()
        return templates.TemplateResponse("organization_documents_add.html", {"request": request, "org": org_fresh, "message": f"Erro ao enviar documento: {str(e)}", "success": False})

@router.post("/organizations/{org_id}/documents/{doc_id}/delete")
def organization_documents_delete(request: Request, org_id: UUID, doc_id: UUID):
    if 'user_id' not in request.session:
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/", status_code=302)
    db = templates.env.globals.get('SessionLocal')() if templates.env.globals.get('SessionLocal') else None
    if db is None:
        from service.orm import SessionLocal, Document, DocumentFile, UserOrganization
        db = SessionLocal()
    from service.orm import Document, DocumentFile, UserOrganization
    user_id = request.session['user_id']
    user_org = db.query(UserOrganization).filter(UserOrganization.user_id == user_id, UserOrganization.organization_id == org_id).first()
    if not user_org:
        db.close()
        from fastapi.responses import RedirectResponse
        return RedirectResponse(url="/organizations", status_code=302)
    doc = db.query(Document).filter(Document.id == doc_id, Document.organization_id == org_id).first()
    if not doc:
        db.close()
        return RedirectResponse(f"/organizations/{org_id}/documents", status_code=303)
    # Remove file from disk
    doc_file = db.query(DocumentFile).filter(DocumentFile.document_id == doc_id).first()
    if doc_file:
        import os
        try:
            if os.path.exists(doc_file.file_path):
                os.remove(doc_file.file_path)
        except Exception:
            pass
        db.delete(doc_file)
    db.delete(doc)
    db.commit()
    db.close()
    from fastapi.responses import RedirectResponse
    return RedirectResponse(f"/organizations/{org_id}/documents", status_code=303)
